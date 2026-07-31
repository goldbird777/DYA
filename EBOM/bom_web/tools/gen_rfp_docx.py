# -*- coding: utf-8 -*-
"""PLM / ERP 도입 RFP — 기준정보 수립 방안 (Word 문서 생성).

이 시스템을 만들며 실측한 내용을 근거로 «기준정보가 어떻게 수립되어야 하는가»를
설명하는 제안요청서다. PLM 업체·ERP 업체에 함께 보낸다.

범위: PLM·ERP 도입과 기준정보 체계에 한정한다. SCM·MES·더존 ERP 연계는 이번
RFP 범위 밖이며 별도 진행한다.

숫자는 전부 실제 검증 결과다. 추정치는 «추정»이라고 명시한다.

생성:  python tools/gen_rfp_docx.py
결과:  data/rfp/DYA_PLM_ERP_RFP.docx
"""
import os
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_DIR = os.path.join(BASE, 'data', 'rfp')
OUT_PATH = os.path.join(OUT_DIR, 'DYA_PLM_ERP_RFP.docx')

NAVY = RGBColor(0x1A, 0x23, 0x7E)
GRAY = RGBColor(0x54, 0x6E, 0x7A)
FONT = '맑은 고딕'


# ── 서식 도우미 ──────────────────────────────────────────────────────────────
def _set_font(run, size=10.5, bold=False, color=None, name=FONT):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    # 한글 글꼴은 eastAsia 속성을 따로 지정해야 적용된다
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def h1(doc, text):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    _set_font(p.add_run(text), 18, True, NAVY)
    _bottom_border(p)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(5)
    _set_font(p.add_run(text), 13, True, NAVY)
    return p


def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(3)
    _set_font(p.add_run(text), 11, True, RGBColor(0x0D, 0x47, 0xA1))
    return p


def para(doc, text, size=10.5, bold=False, color=None, indent=0, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.4
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    _set_font(p.add_run(text), size, bold, color)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.35
    _set_font(p.add_run(text), 10.5)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.35
    _set_font(p.add_run(text), 10.5)
    return p


def _bottom_border(p):
    pPr = p._element.get_or_add_pPr()
    bd = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:color'), '1A237E')
    bd.append(bottom)
    pPr.append(bd)


def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def table(doc, headers, rows, widths=None, note=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p.add_run(str(htxt)), 9.5, True, RGBColor(0xFF, 0xFF, 0xFF))
        _shade(hdr[i], '1A237E')
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            _set_font(p.add_run(str(v)), 9.5)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                if i < len(r.cells):
                    r.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    if note:
        para(doc, note, 9, color=GRAY, after=8)
    return t


def box(doc, title, lines, fill='F1F6FE'):
    """강조 박스 — 1칸짜리 표로 만든다."""
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    c = t.rows[0].cells[0]
    _shade(c, fill)
    c.text = ''
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    _set_font(p.add_run(title), 10.5, True, NAVY)
    for ln in lines:
        pp = c.add_paragraph()
        pp.paragraph_format.space_after = Pt(2)
        pp.paragraph_format.line_spacing = 1.35
        _set_font(pp.add_run(ln), 10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ── 문서 본문 ────────────────────────────────────────────────────────────────
def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)      # A4
    sec.left_margin = sec.right_margin = Cm(2.2)
    sec.top_margin = sec.bottom_margin = Cm(2.0)

    st = doc.styles['Normal']
    st.font.name = FONT
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

    # ── 표지 ─────────────────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run('제안요청서 (RFP)'), 15, True, GRAY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    _set_font(p.add_run('PLM / ERP 도입'), 30, True, NAVY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run('기준정보(Master Data) 수립 방안'), 17, True, NAVY)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run('― 자동차 시트 E-BOM / M-BOM 기준정보 체계 ―'), 11, False, GRAY)
    for _ in range(7):
        doc.add_paragraph()
    for line, sz, bold in (('주식회사 대유에이피', 13, True),
                           ('연구소 선행파트', 11, False),
                           ('2026', 11, False)):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        _set_font(p.add_run(line), sz, bold)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run('본 문서는 제안 목적으로 제공되며, 수신처 외 배포를 금합니다.'), 9, False, GRAY)

    # ══ 1. 문서 개요 ══════════════════════════════════════════════
    h1(doc, '1. 문서 개요')

    h2(doc, '1.1 목적')
    para(doc, '본 문서는 자동차 시트 제품의 설계·생산 기준정보(Master Data)를 체계화하기 위한 '
              'PLM 및 ERP 시스템 도입 제안을 요청하기 위해 작성되었다.')
    para(doc, '일반적인 시스템 도입 제안요청서와 달리, 본 문서는 발주사가 기준정보 문제를 '
              '직접 분석하고 일부를 자체 구축하여 검증한 결과를 근거로 한다. 따라서 요구사항이 '
              '추상적 기능 나열이 아니라 실제 데이터로 확인된 문제와 그 해결 원리를 기준으로 기술된다.')
    box(doc, '본 문서의 특징', [
        '· 모든 수치는 실제 BOM 파일을 분석한 실측값이다. 추정치는 «추정»으로 명시한다.',
        '· 발주사가 이미 자체 구축·검증한 기능은 «구축완료»로 표시하며, 제안 범위에서 제외한다.',
        '· 기준정보 수립 원칙을 먼저 제시하고, 그 원칙을 만족하는지를 평가 기준으로 삼는다.',
    ])

    h2(doc, '1.2 제출 대상')
    table(doc,
          ['구분', '대상', '요청 내용'],
          [['PLM', '제품수명주기관리 솔루션 공급사',
            '3D CAD 연동, 도면·품목 관리, 설계 변경 관리, E-BOM 관리'],
           ['ERP', '전사자원관리 솔루션 공급사',
            '입고단위 BOM 전개, 단가·원가 산출, 구매·생산 연계']],
          widths=[2.2, 5.6, 8.6])

    h2(doc, '1.3 범위 및 제외')
    h3(doc, '포함 범위')
    for t in ('설계 기준정보(E-BOM) 및 생산 기준정보(M-BOM) 체계 수립',
              '품목 마스터(품번·품명·재질·중량·도면) 관리',
              '사양 코드 체계(PEL CODE)를 기준으로 한 설계·생산 용어 통합',
              '3D CAD 데이터와 품번의 연결 방식',
              'BOM 기준정보 검증 및 변경 관리'):
        bullet(doc, t)

    h3(doc, '제외 범위 (이번 RFP에서 다루지 않음)')
    para(doc, '아래 항목은 별도 과제로 진행하며 본 RFP의 평가 대상이 아니다. 다만 향후 연계 '
              '가능성은 제안서에 기술해 주기 바란다.')
    for t in ('SCM(공급망관리) 시스템 구축',
              'MES(제조실행시스템) 구축 및 설비 연동',
              '더존 ERP 등 기존 회계·인사 시스템과의 통합',
              '생산 실적 집계 및 원가 정산 자동화'):
        bullet(doc, t)

    h2(doc, '1.4 용어')
    table(doc,
          ['용어', '설명'],
          [['E-BOM', '설계 BOM. 설계 부서가 관리하는 제품 구성.'],
           ['M-BOM', '생산 BOM. 생산관리 부서가 관리하며 HKMC ALC 코드집 기준.'],
           ['PEL CODE', 'HKMC가 부여하는 6자리 사양 코드. 본 체계의 기준 키.'],
           ['VC', 'Variant Configuration. 사양 조합 단위. 차종당 수십 개.'],
           ['ALC 코드집', 'HKMC가 좌석 열별로 제공하는 사양·품번 대응표.'],
           ['Q파트 종합', 'HKMC 생산계획표. VC별 생산 수량과 사양 조합 포함.'],
           ['BRE', 'BOM Report Excel. 고객사가 제공하는 공장별 1레벨 품번 자료.'],
           ['배타군', '한 제품이 값을 하나만 가지는 사양 축(예: 원단, 에어백).']],
          widths=[3.2, 13.2])

    # ══ 2. 업무 현황 ══════════════════════════════════════════════
    h1(doc, '2. 업무 현황')

    h2(doc, '2.1 제품 및 고객')
    para(doc, '당사는 자동차 시트를 설계·생산하며 주 고객은 HKMC(현대·기아)이다. '
              '시트는 좌석 열(1열 운전석/조수석, 2열 좌·중·우)별로 구성되며, '
              '각 열마다 원단·에어백·열선·통풍·헤드레스트·럼버서포트 등 사양 조합이 존재한다.')
    para(doc, '하나의 차종에 대해 사양 조합(VC)은 수십 개이며, 각 VC마다 1레벨 완제품 품번이 '
              '부여된다. 그 아래로 2~7레벨의 하위 부품이 전개된다.')

    h2(doc, '2.2 문서 흐름')
    para(doc, '기준정보는 아래 흐름으로 만들어진다. 각 단계에서 서로 다른 부서가 서로 다른 '
              '양식을 사용하는 것이 문제의 출발점이다.')
    table(doc,
          ['단계', '입력 문서', '주관', '산출물'],
          [['1', 'HKMC PEL(부품사양서)', '설계', 'VC별 사양 목록'],
           ['2', 'BRE(공장별 1레벨 품번)', '설계', 'VC ↔ 1레벨 품번 대응'],
           ['3', '위 1~2 종합', '설계', 'E-BOM (표준 BOM 양식)'],
           ['4', 'Q파트 종합 + ALC 코드집 5종', '생산관리', 'M-BOM / DYA ALC-2 코드'],
           ['5', 'E-BOM ↔ M-BOM 대조', '설계·생관', '기준정보 정합성 확인']],
          widths=[1.5, 5.6, 2.4, 6.9])

    h2(doc, '2.3 현재의 작업 방식과 한계')
    para(doc, '현재 BOM 작성은 대부분 엑셀 수작업이다. 실제 파일 기준 규모는 다음과 같다.')
    table(doc,
          ['항목', '실측값', '비고'],
          [['부품 행 수', '452행', '1열 운전석 LHD 1개 파일 기준'],
           ['VC 열 수', '53열', '사양 조합 수'],
           ['수량 매트릭스 크기', '23,956셀', '452 × 53'],
           ['실제 값이 채워진 셀', '6,009셀', '나머지는 미해당'],
           ['좌석 위치 수', '5개', 'FRT LH/RH, RR BACK LH/CUSH/RH'],
           ['전 좌석 환산 매트릭스', '약 120,000셀', '추정 (23,956 × 5)']],
          widths=[5.4, 3.6, 7.4],
          note='※ 출처: 250425 NQ5 PE FRT SEAT BOM (운전석 LHD) 파일 분석 결과.')

    para(doc, '이 규모를 사람이 손으로 채우기 때문에 다음 문제가 발생한다.')
    for t in ('작성에 1~2일이 소요되며 그동안 다른 업무를 병행하기 어렵다.',
              '사양이 변경되면 어느 셀이 영향을 받는지 추적할 수 없어 처음부터 다시 작성한다.',
              '설계와 생산관리가 각자 문서를 유지하여 두 문서의 불일치를 사후에 발견한다.',
              '품번·수량 오기입을 사람이 눈으로 검토하므로 누락이 발생한다.'):
        bullet(doc, t)

    # ══ 3. 기준정보 문제 정의 ══════════════════════════════════════
    h1(doc, '3. 기준정보 문제 정의')
    para(doc, '본 장은 본 RFP의 핵심이다. 아래 네 가지 문제는 모두 실제 데이터 분석으로 확인된 '
              '것이며, 제안하는 솔루션이 이 문제들을 어떻게 해결하는지를 평가 기준으로 삼는다.')

    h2(doc, '3.1 문제 ① — 같은 사양이 문서마다 다른 이름으로 적힌다')
    para(doc, '가장 근본적인 문제다. 동일한 사양이 문서와 부서에 따라 서로 다른 표기를 갖는다.')
    table(doc,
          ['사양', 'PEL 코드 마스터', 'BOM N열(설계)', 'VC 헤더(생관)'],
          [['인조가죽', 'A/LEA(1), A/LEA(2)', 'A/LEATHER', '인조'],
           ['PU 인조가죽', 'PU A/LEATHER', 'PU', 'PU'],
           ['천연가죽', 'P/LEA', 'P/LEATHER', '천연'],
           ['직물', 'CLOTH1, CLOTH2', 'CLOTH', '천'],
           ['사이드 에어백', 'THORAX / T & P', 'THORAX / T&P', 'T / T&P'],
           ['시트 열선', 'HTR(2STEP), HTR(3STEP)', 'S/HTR, V/HTR', '히터'],
           ['통풍', 'VENT', 'VENT', '통풍'],
           ['백보드', 'B/BOARD', 'B/COVER', '―']],
          widths=[3.0, 4.6, 4.4, 4.4])

    box(doc, '이 문제가 만드는 결과', [
        '· 시스템이 두 문서를 자동 대조할 수 없다. 사람이 눈으로 맞춰야 한다.',
        '· 신규 사양이 생길 때마다 프로그램 코드를 고쳐야 한다.',
        '· 실측: 설계 BOM의 사양 표기 중 PEL 마스터로 해석되지 않는 항목이 9종 존재했다.',
    ], fill='FFEBEE')

    h2(doc, '3.2 문제 ② — 사양은 품명이 아니라 품번에 들어 있다')
    para(doc, 'BOM의 품명은 구조 명칭이며 사양을 담지 않는다. 사양 차이는 오직 품번으로만 '
              '구분된다. 이는 품명 기반 검색·검증이 원리적으로 불가능함을 뜻한다.')
    table(doc,
          ['항목', '실측값'],
          [['2레벨 이하 부품 행 수', '367행'],
           ['고유 품번 수', '359개'],
           ['고유 품명 수', '169개'],
           ['같은 품명에 여러 품번이 달린 «변형 그룹»', '42개 그룹 / 240행 (전체의 65.4%)'],
           ['PEL 사양 65개 중 품명에 등장하는 것', '5개 (VENT, T&P, PWR, IMS, L/SUPT)']],
          widths=[9.4, 7.0])
    para(doc, '예를 들어 품명 «BACK ASSY-FR SEAT, LH» 하나에 품번이 47종 존재한다. '
              '«CUSH ASSY-FR SEAT, LH»는 14종이다. 모두 품명은 동일하고 사양만 다르다.')
    para(doc, '실제 사양은 BOM의 DESCRIPTION 열(N열)에 «THORAX+CLOTH+L/SUPT» 형태로 '
              '기록되어 있다. 즉 기준정보 체계는 품명이 아니라 이 사양 문자열과 품번을 '
              '연결하는 방식으로 설계되어야 한다.')

    h2(doc, '3.3 문제 ③ — 사양 조합을 단순 비교하면 전건이 불일치로 나온다')
    para(doc, '설계 BOM과 생산 ALC 코드집의 사양을 집합으로 비교한 결과, 완전일치율이 '
              '0%였다. 그러나 실제 오류는 아니었다. 원인은 세 가지다.')
    table(doc,
          ['원인', '설명', '예'],
          [['계열 vs 특정값', '설계는 상위 개념, 생관은 하위 특정값을 기재',
            'A/LEATHER ↔ A/LEA(1)'],
           ['표기 세밀도 차이', '설계는 일반명, 생관은 세부 코드',
            'USB ↔ USB(27W)\nL/SUPT ↔ L/SUPT(4CELL)'],
           ['관리 축 자체가 다름', '한쪽만 관리하는 사양 축이 존재',
            '헤드레스트, 백보드']],
          widths=[3.4, 7.4, 5.6])
    box(doc, '시사점', [
        '기준정보 검증은 «완전일치»가 아니라 «축(배타군) 단위 비교»로 설계되어야 한다.',
        '또한 계층(상위 개념이 하위를 포괄) 판정이 반드시 필요하다.',
        '이 원리를 적용한 결과 오탐 116건이 0건으로 감소했다(실측).',
    ])

    h2(doc, '3.4 문제 ④ — 사양이 바뀌면 전체를 다시 만든다')
    para(doc, '현재는 PEL 사양이 변경되면 영향 범위를 추적할 수 없어 BOM 전체를 재작성한다. '
              '변경 이력이 남지 않으므로 «무엇이 왜 바뀌었는가»를 사후에 확인할 수 없다.')
    para(doc, '참고로 당사가 검토한 외부 BOM 작성 프로그램(외주 개발, 웹 기반)의 소스를 '
              '분석한 결과, 변경 재적용·차이 비교에 해당하는 기능이 구현되어 있지 않았다. '
              '해당 프로그램은 레벨·수량을 웹 화면에서 수기 입력하는 구조로, '
              '사양 변경 시 재작업 문제를 해결하지 못한다.')
    box(doc, '요구 원칙', [
        '기준정보 시스템은 «결과값»이 아니라 «조건»을 저장해야 한다.',
        '결과값을 저장하면 입력이 바뀔 때 다시 만들어야 하지만,',
        '조건을 저장하면 입력이 바뀌어도 재평가만 하면 된다.',
    ])

    # ══ 4. 기준정보 수립 원칙 ══════════════════════════════════════
    h1(doc, '4. 기준정보 수립 원칙')
    para(doc, '3장의 문제를 해결하기 위해 당사가 도출하고 자체 시스템으로 검증한 원칙이다. '
              '제안 솔루션이 아래 원칙을 어떻게 구현하는지 기술해 주기 바란다.')

    h2(doc, '원칙 1 — PEL CODE를 단일 기준 키로 삼는다')
    para(doc, '설계·생산·품질이 각자 용어를 쓰는 한 통합은 불가능하다. HKMC가 부여하는 '
              'PEL CODE(6자리)를 유일한 기준으로 두고, 모든 부서 용어를 여기에 매핑한다.')
    for t in ('PEL CODE는 고객사가 부여하므로 사내 사정으로 바뀌지 않는다.',
              'ALC 코드집이 이미 PEL CODE를 사용하므로 생산 측 연결이 자연스럽다.',
              '설계 BOM의 사양 문자열만 PEL CODE로 해석하면 양측이 연결된다.'):
        bullet(doc, t)

    h2(doc, '원칙 2 — 용어 사전을 마스터 데이터로 관리한다')
    para(doc, '각 PEL CODE에 «별칭» 목록을 부여한다. 설계 표기·생관 표기·한글 표기를 모두 '
              '별칭으로 등록하면 시스템이 동일 사양으로 인식한다.')
    table(doc,
          ['PEL 코드', '표준 사양', '등록된 별칭'],
          [['8811F2', 'A/LEA(1)', 'ARTIFICIAL LEATHER - ONE TONE, A/LEATHER, 인조'],
           ['8811M1', 'PU A/LEATHER', 'PU ARTIFICIAL LEATHER - ONE TONE, PU, A/LEATHER'],
           ['8811E2', 'P/LEA', 'PURE LEATHER, P/LEATHER, 천연'],
           ['8811C1', 'CLOTH1', 'CLOTH, 천']],
          widths=[2.4, 3.4, 10.6])
    box(doc, '핵심 효과', [
        '용어가 추가되어도 프로그램 코드를 고치지 않는다. 마스터에 한 줄 추가하면 된다.',
        '실측: 별칭 5줄을 등록하자 원단 축 대조가 49건 전부 «일치»로 판정되었다.',
    ])
    para(doc, '주의: 별칭이 다른 사양의 표준명과 충돌하면 오인이 발생한다. 실측에서 '
              '럼버서포트(L/SUPT)의 설명란에 «PWR»이 적혀 있어, 이를 그대로 별칭으로 쓰면 '
              '모든 PWR 표기를 럼버로 오인하는 문제가 확인되었다. 시스템은 이러한 충돌을 '
              '자동으로 배제해야 한다.')

    h2(doc, '원칙 3 — 사양을 «배타군(축)»으로 묶는다')
    para(doc, '사양은 종류별로 묶이며, 각 묶음에서 값은 하나만 선택된다. 이 구조를 명시해야 '
              '검증이 가능하다.')
    table(doc,
          ['배타군(축)', '값 예시', '택1 여부'],
          [['원단', 'CLOTH / A/LEATHER / PU / P/LEATHER', '예'],
           ['에어백', 'THORAX / T & P', '예'],
           ['헤드레스트', 'H_UP/DN / H_SLID\'G / 4WAY H/REST', '예'],
           ['럼버서포트', 'L/SUPT / L/SUPT(2CELL) / L/SUPT(4CELL)', '예'],
           ['통풍', 'VENT', '―'],
           ['기타 옵션', 'USB / SBR / TABLE / B/BOARD 등', '아니오 (공존 가능)']],
          widths=[3.2, 8.8, 4.4])
    box(doc, '반드시 주의할 점', [
        '모든 사양을 배타군으로 취급하면 안 된다. USB·SBR·TABLE처럼 공존 가능한 항목까지',
        '택1로 보면 정상 조합이 전부 오류로 판정된다.',
        '실측: 이 구분 없이 검증했을 때 116건이 전건 오탐이었고, 배타군을 한정하자 0건이 되었다.',
    ], fill='FFF8E1')

    h2(doc, '원칙 4 — 계층(상위·하위) 관계를 표현한다')
    para(doc, '사양에는 포함 관계가 존재한다. 상위 개념으로 기재된 부품은 하위 특정값에도 '
              '적용되어야 한다.')
    for t in ('A/LEATHER(인조가죽 계열) ⊃ A/LEA(1), A/LEA(2), PU A/LEATHER',
              'PWR(IMS)는 PWR을 포함한다 — IMS 시트는 파워시트이므로',
              'L/SUPT(4CELL)은 L/SUPT를 포함한다'):
        bullet(doc, t)
    para(doc, '계층 범위는 프로그램이 아니라 마스터 데이터가 결정해야 한다. 실측 검증에서 '
              'PU를 인조가죽 계열로 등록했을 때와 별개 원단으로 등록했을 때 검출 건수가 '
              '14건과 2건으로 달라졌다. 즉 계층 정의가 곧 업무 규칙이며, 이를 '
              '사용자가 직접 관리할 수 있어야 한다.')

    h2(doc, '원칙 5 — 검증은 «필요조건»만 판정한다')
    para(doc, '사양이 맞다고 해서 반드시 그 VC에 배정되는 것은 아니다. 따라서 검증은 '
              '«어긋난 것»만 잡고, «없는 것»은 오류로 단정하지 않아야 한다.')
    table(doc,
          ['판정', '조건', '조치 위치'],
          [['일치', '양쪽 축 값이 겹침', '―'],
           ['불일치', '양쪽 값이 있는데 서로 다름', 'BOM 파일 또는 ALC 확인'],
           ['보완필요(마스터)', '용어는 있으나 마스터에 미등록', 'PEL 마스터 설명란에 별칭 추가'],
           ['보완필요(문서)', '해당 축이 문서에 아예 없음', 'BOM 파일 확인']],
          widths=[3.4, 6.4, 6.6])
    box(doc, '중요', [
        '오류 메시지는 «무엇이 틀렸다»가 아니라 «어디를 고쳐야 한다»를 알려주어야 한다.',
        '특히 «마스터를 고쳐라»와 «문서를 고쳐라»는 담당자가 다르므로 반드시 구분되어야 한다.',
    ])

    h2(doc, '원칙 6 — 원본 서식을 보존한다')
    para(doc, 'BOM은 협력사와 고객사에 배포되는 문서이므로 양식이 유지되어야 한다. '
              '시스템이 데이터를 다시 조립해 새 파일을 만들면 병합·색상·열너비가 사라진다.')
    para(doc, '당사 자체 구축 시스템은 원본 워크북을 열어 «변경된 셀만 덮어쓰는» 방식으로 '
              '이를 해결했다. 실측 결과 병합 셀 29개, 열너비, 셀 배경색이 원본과 동일하게 '
              '유지되었다.')

    # ══ 5. 자체 구축 현황 ══════════════════════════════════════════
    h1(doc, '5. 자체 구축 현황 (제안 범위 제외)')
    para(doc, '당사는 위 원칙을 검증하기 위해 아래 기능을 자체 구축하여 운영 중이다. '
              '해당 항목은 제안 범위에서 제외하며, 견적에 포함하지 않기 바란다.')
    para(doc, '단, 도입 솔루션으로 이관하는 것이 타당하다고 판단되는 항목이 있다면 그 근거와 '
              '함께 제안할 수 있다. 이 경우 이관 범위와 비용을 별도 항목으로 명시해야 한다.')

    h2(doc, '5.1 구축 항목 요약')
    table(doc,
          ['No', '기능', '핵심 검증 결과'],
          [['1', 'BOM 자동 검증 (오류 6종)', '실제 오류 검출, 오탐 0건'],
           ['2', 'PEL 코드 기반 사양 체계', '코드 수정 없이 사양 확장'],
           ['3', 'E-BOM ↔ M-BOM 기준정보 대조', '축 단위 판정, 보완 대상 자동 도출'],
           ['4', '웹 BOM 편집기 (서식 보존)', '렌더 1.29초 / 36,711셀'],
           ['5', '품목 마스터 (2D/3D)', 'BOM 업로드 시 409건 자동 등록']],
          widths=[1.2, 6.0, 9.2])

    h2(doc, '5.2 BOM 자동 검증')
    para(doc, '업로드된 BOM에서 6종의 오류를 자동 검출하고, 각 오류마다 수정 위치를 지목한다.')
    table(doc,
          ['오류 유형', '내용'],
          [['오사양 누락', '1레벨 사양이 하위 부품에 없음'],
           ['사양 불일치', '하위 부품 사양이 1레벨에 없음'],
           ['수량 오기입', '소수점·음수·비정상 대수량'],
           ['레벨 중복', '같은 VC에 동일 ASSY 중복 배정'],
           ['P/NO 중복', '동일 ASSY 하위 품번 중복'],
           ['사양 배정 모순', '부품 사양이 배정된 VC의 사양과 충돌']],
          widths=[3.6, 12.8])
    para(doc, '실측 사례: 품번 88751-P1500이 «SLAB SPONGE-FR H/REST»와 «PAPER TAPE» '
              '두 부품에 동일하게 등재된 오류를 검출하였다.', 10, color=GRAY)

    h2(doc, '5.3 PEL 코드 기반 사양 체계')
    para(doc, '4장의 원칙 1~4를 구현한 것으로, 마스터 데이터만으로 사양 체계가 확장된다.')
    for t in ('PEL 마스터의 설명란을 용어 사전으로 사용 (쉼표로 별칭 등록)',
              '옵션그룹을 배타군으로 사용',
              '계열/계층 자동 판정',
              '충돌 별칭 자동 배제'):
        bullet(doc, t)

    h2(doc, '5.4 E-BOM ↔ M-BOM 기준정보 대조')
    para(doc, '설계 BOM과 생산 ALC 코드집을 축 단위로 대조하여 정합성을 확인한다.')
    table(doc,
          ['축', '판정 결과', '해석'],
          [['원단', '49건 전부 일치', '정합 완료'],
           ['에어백', '45건 일치', '정합 완료'],
           ['통풍', '13건 일치', '정합 완료'],
           ['헤드레스트', '49건 보완 필요', '마스터에 용어 미등록'],
           ['럼버서포트', '32건 값 상이', '표기 세밀도 차이']],
          widths=[3.0, 5.0, 8.4],
          note='※ 출처: NQ5 E-BOM(1열 운전석) × M-BOM ALC 코드집 대조 결과. 품번 매칭 49건 기준.')
    para(doc, '이 결과가 1레벨 BOM 자동 생성의 선행 자료가 된다. 즉 «무엇을 표준화해야 '
              '자동화가 가능한가»를 데이터로 제시한다.')

    h2(doc, '5.5 웹 BOM 편집기')
    para(doc, '엑셀 BOM을 웹에서 열람·편집하고, 원본 서식을 유지한 채 다운로드한다.')
    table(doc,
          ['항목', '실측값'],
          [['처리 규모', '485행 × 75열 (셀 9,392개)'],
           ['화면 렌더 시간', '1.29초 / 36,711셀'],
           ['서식 재현', '열너비·행높이·색상·병합 원본과 동일'],
           ['다운로드 서식 보존', '병합 29개, 열너비, 셀 배경색 일치 확인'],
           ['편집 잠금', '동시 1명, 30분 무활동 자동 해제'],
           ['변경 이력', '저장 시 리비전 증가, 변경 셀 단위 기록']],
          widths=[5.4, 11.0])

    h2(doc, '5.6 품목 마스터')
    para(doc, 'BOM 업로드 시 전 레벨 품번·품명이 자동 등록되며, 품목별 스펙과 도면을 관리한다.')
    for t in ('자동 등록: BOM 1회 업로드로 409건 등록 (실측)',
              '스펙 20개 필드: OEM·고객품번·재질·MS SPEC·카티아 중량·표면처리 등',
              '도면(pdf/dwg/dxf/이미지) 및 첨부파일 관리',
              '저장 시 REVISION 자동 증가 및 이력 보관',
              '자동 등록은 빈 칸만 채우며 수기 입력 스펙은 보존'):
        bullet(doc, t)

    # ══ 6. PLM 요구사항 ═══════════════════════════════════════════
    h1(doc, '6. PLM 요구사항')
    para(doc, '요구 수준은 «필수」와 «선택»으로 구분한다. 필수 항목을 충족하지 못하는 경우 '
              '그 사유와 대안을 제안서에 명시해야 한다.')

    h2(doc, '6.1 기준정보 관리')
    table(doc,
          ['No', '요구사항', '수준'],
          [['P-01', 'PEL CODE를 기준 키로 하는 사양 마스터 관리', '필수'],
           ['P-02', '사양별 별칭(동의어) 등록 및 자동 인식', '필수'],
           ['P-03', '배타군(택1 사양 축) 정의 및 관리', '필수'],
           ['P-04', '사양 간 계층(포함) 관계 정의', '필수'],
           ['P-05', '마스터 변경만으로 사양 확장 (코드 수정 불요)', '필수'],
           ['P-06', '충돌 별칭 자동 검출 및 배제', '선택'],
           ['P-07', '사양 마스터 변경 이력 및 승인 절차', '필수']],
          widths=[1.6, 11.0, 3.8])

    h2(doc, '6.2 BOM 관리')
    table(doc,
          ['No', '요구사항', '수준'],
          [['P-11', '다단(1~7레벨) BOM 구조 관리', '필수'],
           ['P-12', 'VC(사양 조합)별 BOM 전개', '필수'],
           ['P-13', '사양 조건에 따른 부품 자동 배정', '필수'],
           ['P-14', '엑셀 업로드·다운로드 시 원본 서식 유지', '필수'],
           ['P-15', 'BOM 변경 시 영향 범위 자동 분석', '필수'],
           ['P-16', '변경분만 재생성 (전체 재작성 불요)', '필수'],
           ['P-17', 'BOM 버전·리비전 관리 및 비교', '필수'],
           ['P-18', '동시 편집 제어 (잠금 또는 병합)', '필수']],
          widths=[1.6, 11.0, 3.8])
    box(doc, 'P-14 · P-16 은 타협 불가 항목', [
        'P-14: BOM은 협력사·고객사 배포 문서이므로 양식이 유지되어야 한다.',
        'P-16: 사양 변경 시 전체 재작성은 현재 방식의 가장 큰 문제이며,',
        '        이를 해결하지 못하면 도입 효과가 없다.',
    ], fill='FFEBEE')

    h2(doc, '6.3 3D CAD 연동 — 협의 필요 항목')
    para(doc, '본 항목은 당사 내부에서도 방향을 검토 중이며, 업체 제안을 요청한다.')
    h3(doc, '현황 및 제약')
    for t in ('시트 어셈블리 구성 부품이 200~400개 규모이다.',
              '400개 3D 파일을 개별 관리하는 동종 업체는 확인되지 않았다.',
              '2·3레벨 어셈블리의 카티아 파일을 모두 작성하는 것은 현실적이지 않다.'):
        bullet(doc, t)
    h3(doc, '당사 검토안')
    para(doc, '아래 세 안을 검토 중이며, 각 안의 실현 가능성과 비용에 대한 의견을 요청한다.')
    table(doc,
          ['안', '내용', '검토 의견 요청'],
          [['1안', '단품 도면(BRKT·COVER류, 4레벨)까지 카티아에서 관리하고 품번과 연결. '
                   '2·3레벨은 1레벨 파일을 열람하는 «동일 도면 열람»으로 대체',
            '해당 기능 지원 여부'],
           ['2안', 'Product 구조 업로드 시 부모-자식 E-BOM 자동 생성',
            '실제 운영 사례, 파일 수 대비 비용'],
           ['3안', 'MNL 최저/최고, PWR 최저/최고 4종만 도면 등록. '
                   '해당 1레벨 품번에 표기하여 열람',
            '기준정보 관점의 타당성']],
          widths=[1.4, 9.0, 6.0])
    h3(doc, '공통 확인 요청')
    for t in ('도면 파일 수가 수백 개일 때의 라이선스·스토리지 비용 산정 방식',
              '품번 ↔ 3D 파일 연결 방식 및 품번 변경 시 처리',
              '설계 변경(EO) 발생 시 도면·BOM 동시 개정 처리 방식',
              '카티아 외 CAD 도구 사용 시의 연동 범위'):
        bullet(doc, t)

    h2(doc, '6.4 품목 관리')
    table(doc,
          ['No', '요구사항', '수준'],
          [['P-21', '품목 마스터(품번 유일키) 관리', '필수'],
           ['P-22', 'BOM 등록 시 전 레벨 품번 자동 등록', '필수'],
           ['P-23', '품목별 스펙(재질·중량·MS SPEC·표면처리 등) 관리', '필수'],
           ['P-24', '도면·문서 첨부 및 버전 관리', '필수'],
           ['P-25', '품목 변경 이력 및 승인 절차', '필수'],
           ['P-26', '자동 등록이 수기 입력 데이터를 덮어쓰지 않을 것', '필수']],
          widths=[1.6, 11.0, 3.8])

    # ══ 7. ERP 요구사항 ═══════════════════════════════════════════
    h1(doc, '7. ERP 요구사항')
    para(doc, 'ERP는 설계 기준정보를 받아 구매·생산·원가로 연결하는 역할을 담당한다. '
              '본 RFP에서는 기준정보 연계와 BOM 전개에 한정하여 요구한다.')

    h2(doc, '7.1 입고단위 BOM')
    para(doc, '현재 E-BOM(설계)과 M-BOM(생산)은 자체 구축하였으나, 영업·구매가 활용하는 '
              '입고단위 BOM 전개는 미구현 상태이다.')
    table(doc,
          ['No', '요구사항', '수준'],
          [['E-01', 'E-BOM으로부터 입고단위 BOM 자동 전개', '필수'],
           ['E-02', '전개 규칙(단위 환산·묶음)을 사용자가 설정', '필수'],
           ['E-03', '협력사별 납품 단위 관리', '필수'],
           ['E-04', 'BOM 개정 시 입고단위 BOM 자동 갱신', '필수'],
           ['E-05', '전개 결과 검증 및 차이 보고', '선택']],
          widths=[1.6, 11.0, 3.8])

    h2(doc, '7.2 단가 및 원가')
    table(doc,
          ['No', '요구사항', '수준'],
          [['E-11', '품목별 단가 이력 관리', '필수'],
           ['E-12', 'BOM 기준 제품 원가 자동 산출', '필수'],
           ['E-13', '사양(VC)별 원가 비교', '필수'],
           ['E-14', '단가 변경 시 원가 영향 분석', '필수'],
           ['E-15', '기존 «영업 단가» 데이터 이관', '필수']],
          widths=[1.6, 11.0, 3.8])
    para(doc, '당사는 현재 영업 단가를 별도 게시판으로 관리하고 있으며, 해당 데이터의 '
              '이관 방안을 제안서에 포함해 주기 바란다.', 10, color=GRAY)

    h2(doc, '7.3 구매·생산 연계')
    table(doc,
          ['No', '요구사항', '수준'],
          [['E-21', 'BOM 기준 소요량 산출(MRP)', '필수'],
           ['E-22', '생산계획(Q파트 종합) 연계', '선택'],
           ['E-23', '협력사 발주 및 입고 관리', '필수'],
           ['E-24', '재고 관리', '필수']],
          widths=[1.6, 11.0, 3.8])
    para(doc, '※ 생산계획 연계(E-22)는 향후 MES 도입과 관련되므로 이번 범위에서는 '
              '«연계 가능성»만 확인한다.', 10, color=GRAY)

    # ══ 8. 연계 요구사항 ══════════════════════════════════════════
    h1(doc, '8. 연계 요구사항')

    h2(doc, '8.1 시스템 간 데이터 흐름')
    para(doc, 'PLM·ERP·자체 시스템 간 기준정보 흐름은 아래와 같이 구성되어야 한다. '
              '기준정보의 «단일 진실 원천(Single Source of Truth)»이 어디인지 명확해야 한다.')
    table(doc,
          ['원천', '데이터', '수신', '주기'],
          [['PLM', '품목 마스터, E-BOM, 도면', 'ERP', '변경 시'],
           ['PLM', '사양 마스터(PEL CODE)', '자체 검증 시스템', '변경 시'],
           ['ERP', '단가, 입고단위 BOM', 'PLM(참조)', '변경 시'],
           ['자체 시스템', 'BOM 검증 결과', 'PLM', '검증 시'],
           ['HKMC', 'PEL 부품사양서, ALC 코드집', 'PLM / 자체 시스템', '수시']],
          widths=[3.0, 5.6, 4.4, 3.4])

    h2(doc, '8.2 연계 요구사항')
    table(doc,
          ['No', '요구사항', '수준'],
          [['I-01', '표준 인터페이스(REST API 등) 제공', '필수'],
           ['I-02', '기준정보 변경 시 연계 시스템 자동 통지', '필수'],
           ['I-03', '연계 실패 시 재처리 및 이력 관리', '필수'],
           ['I-04', '자체 구축 시스템(Python/FastAPI)과의 연동 가능 여부', '필수'],
           ['I-05', '엑셀 일괄 업로드·다운로드 지원', '필수']],
          widths=[1.6, 11.0, 3.8])

    h2(doc, '8.3 향후 확장 (참고)')
    para(doc, '아래는 이번 RFP 범위가 아니나, 향후 도입을 검토 중이므로 확장 가능성을 '
              '제안서에 기술해 주기 바란다.')
    for t in ('MES 연동 — 생산 실적, 설비 데이터',
              'SCM 연동 — 협력사 납기·재고 가시성',
              '기존 회계 시스템(더존 ERP)과의 통합 또는 병행 운영'):
        bullet(doc, t)

    # ══ 9. 도입 범위 및 제외 ══════════════════════════════════════
    h1(doc, '9. 도입 범위 및 제외')

    h2(doc, '9.1 견적 제외 요청 항목')
    para(doc, '5장에 기술한 자체 구축 항목은 이미 운영 중이므로 견적에서 제외하기 바란다. '
              '제안서에는 해당 기능을 «기 보유»로 표기하고 금액을 산정하지 않는다.')
    table(doc,
          ['항목', '상태', '비고'],
          [['BOM 자동 검증 (오류 6종)', '구축 완료', '운영 중'],
           ['PEL 코드 기반 사양 체계', '구축 완료', '운영 중'],
           ['E-BOM ↔ M-BOM 대조', '구축 완료', '운영 중'],
           ['웹 BOM 편집기 (서식 보존)', '구축 완료', '운영 중'],
           ['품목 마스터 (2D/3D)', '구축 완료', '운영 중']],
          widths=[7.0, 3.0, 6.4])
    box(doc, '단, 아래는 별도 협의 대상이다', [
        '· 자체 구축분의 운영·유지보수를 도입 솔루션으로 이관하는 범위',
        '· 이관 시 데이터 마이그레이션 방식과 기간',
        '· 자체 시스템을 병행 운영할 경우의 인터페이스 개발 범위',
    ], fill='FFF8E1')

    h2(doc, '9.2 단계별 도입 계획 (안)')
    table(doc,
          ['단계', '내용', '비고'],
          [['1단계', '기준정보 체계 수립 — 사양 마스터, 품목 마스터 이관',
            '자체 구축 데이터 활용'],
           ['2단계', 'PLM 도입 — E-BOM, 도면, 설계 변경 관리', ''],
           ['3단계', 'ERP 연계 — 입고단위 BOM, 단가·원가', ''],
           ['4단계', '확장 — MES·SCM 연계 검토', '별도 과제']],
          widths=[2.0, 9.4, 5.0],
          note='※ 단계 구분과 기간은 제안사가 조정하여 제시할 수 있다.')

    h2(doc, '9.3 적용 대상')
    para(doc, '초기 적용은 1개 차종으로 한정하여 원리를 검증한 후 확대한다. '
              '당사 자체 시스템도 동일한 방식으로 1개 차종(NQ5) 검증을 진행 중이다.')

    # ══ 10. 평가 기준 ═════════════════════════════════════════════
    h1(doc, '10. 평가 기준 및 제출 요청')

    h2(doc, '10.1 평가 기준')
    table(doc,
          ['구분', '평가 항목', '배점'],
          [['기술', '4장 기준정보 수립 원칙 6가지의 구현 방식', '30'],
           ['기술', '6장·7장 필수 요구사항 충족도', '25'],
           ['기술', '3D CAD 연동 방안의 현실성', '15'],
           ['사업', '동종 업계(자동차 부품) 구축 사례', '10'],
           ['사업', '구축 기간 및 조직', '10'],
           ['비용', '견적 금액 (자체 구축분 제외 여부 포함)', '10']],
          widths=[2.0, 11.0, 3.4])
    para(doc, '기술 평가 비중이 70%이며, 그중 기준정보 원칙 구현이 30%로 가장 높다. '
              '단순 기능 보유 여부보다 3장의 문제를 어떻게 해결하는지를 중점 평가한다.', 10, color=GRAY)

    h2(doc, '10.2 제안서 필수 포함 사항')
    for i, t in enumerate((
            '4장 원칙 6가지 각각에 대한 구현 방식 설명',
            '6장·7장 요구사항별 충족 여부 (충족/부분충족/미충족) 및 사유',
            '6.3 3D CAD 연동 3개 안에 대한 검토 의견',
            '자체 구축 항목의 이관 필요성 판단 및 근거',
            '동종 업계 구축 사례 (가능한 범위에서)',
            '단계별 구축 계획 및 기간',
            '견적 (자체 구축분 제외 명시)',
            '연계 인터페이스 규격'), 1):
        numbered(doc, t)

    h2(doc, '10.3 시연 요청')
    para(doc, '제안 설명 시 아래 항목의 실제 동작을 시연해 주기 바란다. '
              '화면 캡처나 자료가 아닌 실제 시스템 조작을 요청한다.')
    for t in ('사양 마스터에 신규 사양을 추가하고, 코드 수정 없이 BOM에 반영되는 과정',
              'BOM 변경 시 영향 범위가 자동으로 산출되는 과정',
              '엑셀 다운로드 후 원본 서식이 유지되는지 확인',
              '3D CAD 구조를 업로드하여 BOM이 생성되는 과정'):
        bullet(doc, t)

    h2(doc, '10.4 문의')
    para(doc, '본 문서의 기술적 내용에 대한 문의는 아래로 연락 바란다.')
    table(doc,
          ['구분', '내용'],
          [['주관 부서', '주식회사 대유에이피 연구소 선행파트'],
           ['문의 사항', '기준정보 체계, 자체 구축 시스템 관련'],
           ['제출 방식', '별도 안내']],
          widths=[3.4, 13.0])

    # ══ 부록 ═════════════════════════════════════════════════════
    h1(doc, '부록 A. 실측 데이터 요약')
    para(doc, '본 문서에 인용된 수치의 출처와 산출 근거이다.')

    h2(doc, 'A.1 분석 대상 파일')
    table(doc,
          ['구분', '파일'],
          [['E-BOM', '250425 NQ5 PE FRT SEAT BOM — 운전석 파워/천/백커버 신규 사양 추가 (FRT LH)'],
           ['M-BOM', 'GY FRT LH / FRT RH / RR BACK LH / RR CUSH / RR BACK RH ALC 코드집'],
           ['생산계획', 'GY 종합 (Q파트 종합)'],
           ['사양 마스터', 'PEL CODE 마스터 (89개 코드)']],
          widths=[3.0, 13.4])

    h2(doc, 'A.2 BOM 구조')
    table(doc,
          ['항목', '값'],
          [['전체 부품 행', '452행'],
           ['VC(사양 조합) 수', '53개'],
           ['레벨 분포', '1레벨 51 / 2레벨 102 / 3레벨 115 / 4레벨 63 / 5레벨 66 / 6레벨 11 / 7레벨 10'],
           ['수량 매트릭스', '23,956셀 (452 × 53)'],
           ['값이 있는 셀', '6,009셀'],
           ['2레벨 이하 고유 품번', '359개'],
           ['2레벨 이하 고유 품명', '169개'],
           ['변형 그룹', '42개 그룹 / 240행 (65.4%)']],
          widths=[5.4, 11.0])

    h2(doc, 'A.3 사양 대조 결과')
    table(doc,
          ['항목', '값'],
          [['품번 매칭', '49건'],
           ['사양 완전일치율', '0%'],
           ['축 단위 판정 총계', '237건'],
           ['― 일치', '129건'],
           ['― 마스터 보완 필요', '73건'],
           ['― 값 상이', '35건'],
           ['배타군 미한정 시 오탐', '116건 → 한정 후 0건']],
          widths=[5.4, 11.0])

    h2(doc, 'A.4 시스템 처리 성능')
    table(doc,
          ['항목', '값'],
          [['셀 적재', '9,392개 / 10.6초'],
           ['서식 추출', '13.3초 (고유 스타일 141개)'],
           ['화면 렌더', '1.29초 / 36,711셀'],
           ['품목 자동 등록', '409건 / 1회 업로드'],
           ['서식 보존 확인', '병합 29개·열너비·배경색 원본 일치']],
          widths=[5.4, 11.0],
          note='※ 서버 사양: RAM 956MB (클라우드 최소 사양). 확대 시 증설 예정.')

    h1(doc, '부록 B. 기준정보 용어 매핑 예시')
    para(doc, '원칙 2(용어 사전)의 실제 적용 예이다. 아래와 같이 등록하면 세 문서의 표기가 '
              '하나의 PEL 코드로 통합된다.')
    table(doc,
          ['PEL 코드', '표준 사양', '옵션그룹(배타군)', '별칭'],
          [['8811B4', 'A/CLOTH', "COVER'G", 'A/CLOTH(CLOTH+ARTIFICIAL LEATHER)'],
           ['8811C1', 'CLOTH1', "COVER'G", 'CLOTH, 천'],
           ['8811C2', 'CLOTH2', "COVER'G", 'CLOTH, 천'],
           ['8811E2', 'P/LEA', "COVER'G", 'PURE LEATHER, P/LEATHER, 천연'],
           ['8811F2', 'A/LEA(1)', "COVER'G", 'ARTIFICIAL LEATHER - ONE TONE, A/LEATHER, 인조'],
           ['8811F3', 'A/LEA(2)', "COVER'G", 'Artificial Leather(TWO TONE), A/LEATHER, 인조'],
           ['8811M1', 'PU A/LEATHER', "COVER'G", 'PU ARTIFICIAL LEATHER - ONE TONE, PU, A/LEATHER'],
           ['5693A1', 'THORAX', 'Airbag', 'SIDE AIR BAG-FR - FR(THORAX)'],
           ['5693A3', 'T & P', 'Airbag', 'SIDE AIR BAG (T&P)'],
           ['8827A1', 'PWR', 'OPTION', 'POWER SEAT'],
           ['8827A3', 'PWR(IMS)', 'OPTION', 'POWER SEAT(IMS), IMS'],
           ['8828A2', 'L/SUPT', 'LUMBAR', '럼버']],
          widths=[2.2, 3.2, 3.2, 7.8])

    h2(doc, 'B.1 미해결 항목')
    para(doc, '아래는 현재 매핑이 확정되지 않은 항목이며, 도입 솔루션에서 어떻게 처리할지 '
              '제안이 필요하다.')
    table(doc,
          ['표기', '출현 빈도', '문제'],
          [['B/COVER', 'BOM 내 71회', '백보드 계열 8개 코드 중 어느 것에 대응하는지 미확정'],
           ['4W STD / 4W HAN\n2W STD / 2W HAN', 'VC 헤더', 'STD·HANGER 구분 축이 PEL 마스터에 부재'],
           ['S/HTR, V/HTR', 'BOM 내 56회', '작성자 편의 표기로 PEL 사양 아님 (검증 제외 처리)'],
           ['MNL', 'BOM 내 47회', 'PEL 코드 없음. «PWR 부재 = MNL» 부정 조건으로 판정']],
          widths=[3.6, 3.0, 9.8])

    box(doc, '이 항목들이 시사하는 것', [
        '기준정보 체계는 «코드가 없는 사양»을 어떻게 다룰지도 정의해야 한다.',
        '· 신규 코드를 발급받을 것인가',
        '· 부정 조건(다른 코드의 부재)으로 판정할 것인가',
        '· 관리 대상에서 제외할 것인가',
        '제안 솔루션이 이 세 경우를 모두 표현할 수 있는지 확인이 필요하다.',
    ])

    os.makedirs(OUT_DIR, exist_ok=True)
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == '__main__':
    p = build()
    print('생성 완료:', p)
    print('크기:', round(os.path.getsize(p) / 1024, 1), 'KB')
